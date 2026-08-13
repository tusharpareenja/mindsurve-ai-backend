"""Extract plain text from uploaded PDF / Word / plain-text files for the study AI."""

from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

MAX_EXTRACT_CHARS = 12_000

_PDF_TYPES = {"application/pdf"}
_DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_LEGACY_DOC_TYPES = {
    "application/msword",
}
_TEXT_TYPES = {"text/plain", "text/csv", "text/markdown"}


def is_extractable_document(filename: str, content_type: str) -> bool:
    ctype = (content_type or "").split(";")[0].strip().lower()
    suffix = Path(filename or "").suffix.lower()
    if ctype in _PDF_TYPES or suffix == ".pdf":
        return True
    if ctype in _DOCX_TYPES or ctype in _LEGACY_DOC_TYPES or suffix in {
        ".docx",
        ".doc",
    }:
        return True
    if ctype in _TEXT_TYPES or suffix in {".txt", ".csv", ".md"}:
        return True
    return False


def extract_document_text(
    *,
    filename: str,
    content_type: str,
    data: bytes,
) -> str | None:
    """Return extracted text (truncated) or None if the file isn't a readable document."""
    if not data:
        return None
    ctype = (content_type or "").split(";")[0].strip().lower()
    suffix = Path(filename or "").suffix.lower()
    try:
        if ctype in _PDF_TYPES or suffix == ".pdf":
            text = _extract_pdf(data)
        elif (
            suffix == ".docx"
            or ctype in _DOCX_TYPES
            or ctype.endswith("wordprocessingml.document")
        ):
            text = _extract_docx(data)
        elif suffix == ".doc" or ctype in _LEGACY_DOC_TYPES:
            # Legacy binary .doc is not supported by python-docx.
            logger.warning(
                "Legacy .doc upload %s is not supported; ask for .docx or PDF",
                filename or "upload",
            )
            return None
        elif ctype in _TEXT_TYPES or suffix in {".txt", ".csv", ".md"}:
            text = _extract_plain(data)
        else:
            return None
    except Exception:
        logger.exception("Failed to extract text from %s", filename or "upload")
        return None

    cleaned = _normalize(text)
    if not cleaned:
        logger.warning(
            "No readable text extracted from %s (%s, %s bytes)",
            filename or "upload",
            ctype or "unknown-type",
            len(data),
        )
        return None
    if len(cleaned) > MAX_EXTRACT_CHARS:
        return cleaned[:MAX_EXTRACT_CHARS].rstrip() + "\n… [truncated]"
    return cleaned


def _normalize(text: str) -> str:
    lines = [line.strip() for line in (text or "").replace("\r\n", "\n").split("\n")]
    collapsed: list[str] = []
    blank = False
    for line in lines:
        if not line:
            if not blank and collapsed:
                collapsed.append("")
            blank = True
            continue
        blank = False
        collapsed.append(line)
    return "\n".join(collapsed).strip()


def _extract_plain(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    """Extract body paragraphs, tables, and header/footer text via python-docx."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc

    # Guard against the wrong PyPI package named "docx" (Python-2 era).
    if not hasattr(Document, "__call__") and not callable(Document):
        raise RuntimeError("Invalid docx package installed; install python-docx")
    try:
        # python-docx Document is a class; the wrong package exposes a different API.
        document = Document(io.BytesIO(data))
    except TypeError as exc:
        raise RuntimeError(
            "Wrong 'docx' package installed (need python-docx). "
            "Run: pip uninstall docx && pip install python-docx"
        ) from exc

    parts: list[str] = []
    for para in document.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    # Headers / footers often hold titles or campaign copy.
    for section in document.sections:
        for container in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            try:
                for para in container.paragraphs:
                    if para.text and para.text.strip():
                        parts.append(para.text)
            except Exception:
                continue

    # Text boxes / floating shapes often hold campaign copy.
    try:
        seen = {p.strip() for p in parts if p.strip()}
        for box in document.element.iter(qn("w:txbxContent")):
            texts = [
                (node.text or "").strip()
                for node in box.iter(qn("w:t"))
                if (node.text or "").strip()
            ]
            joined = " ".join(texts).strip()
            if joined and joined not in seen:
                parts.append(joined)
                seen.add(joined)
    except Exception:
        pass

    return "\n".join(parts)
